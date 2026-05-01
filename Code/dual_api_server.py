import io
import hashlib
import json
import os
import re
from datetime import datetime, timedelta
from pathlib import Path
from typing import List
from urllib.parse import quote

import docx2txt
import numpy as np
import uvicorn
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from openai import AsyncOpenAI
from PIL import Image
from PyPDF2 import PdfReader
from pydantic import BaseModel, Field

# ================= 1. Configuration & Initialization =================
CODE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = CODE_DIR.parent

load_dotenv(PROJECT_ROOT / ".env")
load_dotenv(CODE_DIR / ".env", override=True)

if os.getenv("HF_OFFLINE", "1").lower() in {"1", "true", "yes", "on"}:
    os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
    os.environ.setdefault("HF_DATASETS_OFFLINE", "1")
    os.environ.setdefault("HF_HUB_OFFLINE", "1")

def project_path_from_env(name: str, default: Path) -> Path:
    raw_value = os.getenv(name)
    path = Path(raw_value).expanduser() if raw_value else default
    if not path.is_absolute():
        path = PROJECT_ROOT / path
    return path.resolve()

DB_SAVE_PATH = project_path_from_env("CHROMA_DB_PATH", PROJECT_ROOT / "Model" / "chroma_db")
LOG_FILE_PATH = project_path_from_env("CHAT_LOG_PATH", PROJECT_ROOT / "Log" / "justitia_chat_logs.jsonl")
EVENT_LOG_PATH = project_path_from_env("EVENT_LOG_PATH", PROJECT_ROOT / "Log" / "platform_events.jsonl")
SERVER_HOST = os.getenv("HUXIN_HOST", "0.0.0.0")
SERVER_PORT = int(os.getenv("HUXIN_PORT", "8000"))

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
if not DEEPSEEK_API_KEY:
    raise ValueError("DEEPSEEK_API_KEY not found. Copy .env.example to Code/.env or project .env first.")

ocr_engine = None
ocr_engine_name = None
ocr_init_error = None


def initialize_ocr_engine():
    """Prefer RapidOCR for Chinese documents, then fall back to EasyOCR."""
    global ocr_engine, ocr_engine_name, ocr_init_error

    try:
        from rapidocr import RapidOCR

        ocr_engine = RapidOCR()
        ocr_engine_name = "RapidOCR(PP-OCRv4)"
        ocr_init_error = None
        print("RapidOCR initialized successfully")
        return
    except Exception as e:
        ocr_init_error = f"RapidOCR initialization failed: {e}"
        print(ocr_init_error)

    try:
        import easyocr

        ocr_engine = easyocr.Reader(["ch_sim", "en"])
        ocr_engine_name = "EasyOCR"
        print("EasyOCR initialized successfully")
    except Exception as e:
        ocr_init_error = f"{ocr_init_error}; EasyOCR initialization failed: {e}" if ocr_init_error else f"EasyOCR initialization failed: {e}"
        print(ocr_init_error)


initialize_ocr_engine()


def run_ocr(image_np: np.ndarray) -> tuple[str, float | None]:
    if ocr_engine is None:
        raise RuntimeError(ocr_init_error or "OCR engine is not initialized")

    if ocr_engine_name and ocr_engine_name.startswith("RapidOCR"):
        result = ocr_engine(image_np)
        lines = [line.strip() for line in getattr(result, "txts", ()) if str(line).strip()]
        scores = [float(score) for score in getattr(result, "scores", ()) if score is not None]
        confidence = round(sum(scores) / len(scores), 4) if scores else None
        return "\n".join(lines), confidence

    result = ocr_engine.readtext(image_np, detail=1, paragraph=False)
    lines = []
    scores = []
    for item in result:
        if len(item) >= 2 and str(item[1]).strip():
            lines.append(str(item[1]).strip())
        if len(item) >= 3:
            scores.append(float(item[2]))
    confidence = round(sum(scores) / len(scores), 4) if scores else None
    return "\n".join(lines), confidence

def now_text() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def backend_log(event: str, detail: str = ""):
    suffix = f" | {detail}" if detail else ""
    print(f"[{now_text()}] {event}{suffix}", flush=True)


def append_jsonl(path: Path, payload: dict):
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(json.dumps(payload, ensure_ascii=False) + "\n")


def record_platform_event(event_type: str, payload: dict):
    event = {
        "timestamp": now_text(),
        "event_type": event_type,
        **payload,
    }
    try:
        append_jsonl(EVENT_LOG_PATH, event)
    except Exception as e:
        backend_log("Event Log Error", str(e))

    preview = json.dumps(payload, ensure_ascii=False)
    backend_log(f"EVENT {event_type}", preview[:600])


def save_chat_log(
    query: str,
    reasoning: str,
    answer: str,
    sources: list,
    user_name: str = "王某某",
    phone: str = "133 3107 4710",
    case_profile: dict | None = None,
):
    """Save chat history and reasoning to JSONL format."""
    log_data = {
        "timestamp": now_text(),
        "user_name": user_name,
        "phone": phone,
        "user_query": query,
        "justitia_thought": reasoning,
        "justitia_answer": answer,
        "reference_sources": [s.get("filename", "Unknown File") for s in sources],
        "case_profile": case_profile or build_case_profile(f"{query}\n{answer}", user_name=user_name, phone=phone),
        "status": "AI 已答复",
    }
    
    try:
        append_jsonl(LOG_FILE_PATH, log_data)
        backend_log("PHASE2 AI_CHAT_SAVED", f"{user_name}({phone}) | {query[:120]}")
    except Exception as e:
        backend_log("Log Error", str(e))

print(f"Loading vector model and local legal database from {DB_SAVE_PATH}...")
try:
    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-m3")
    vectordb = Chroma(persist_directory=str(DB_SAVE_PATH), embedding_function=embeddings)
except Exception as e:
    print(f"Database loading failed: {e}")
    raise

print("Initializing DeepSeek-V4 model...")
client = AsyncOpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com"
)

app = FastAPI(title="Justitia Shield Engine", description="Legal Intelligence Engine for Procuratorate")

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def add_private_network_access_header(request, call_next):
    response = await call_next(request)
    response.headers["Access-Control-Allow-Private-Network"] = "true"
    return response


@app.get("/api/health")
async def health_check():
    return {
        "status": "ok",
        "database_path": str(DB_SAVE_PATH),
        "database_exists": DB_SAVE_PATH.exists(),
        "ocr_ready": ocr_engine is not None,
        "ocr_engine": ocr_engine_name,
        "ocr_init_error": ocr_init_error if ocr_engine is None else None,
    }

# ================= 2. Data Models =================
class ChatRequest(BaseModel):
    query: str
    stream: bool = True
    history: list = Field(default_factory=list)
    top_k: int = 3
    score_threshold: float = 1.2
    mode: str = "spark"
    user_name: str = "王某某"
    phone: str = "133 3107 4710"

class SourceItem(BaseModel):
    filename: str
    score: float
    content_preview: str

class ChatResponse(BaseModel):
    answer: str
    sources: List[SourceItem]

class DocExportRequest(BaseModel):
    title: str = "护薪法律文书"
    html: str
    user_name: str = "王某某"
    phone: str = "133 3107 4710"


class HumanSupportRequest(BaseModel):
    phase: str
    user_name: str = "王某某"
    phone: str = "133 3107 4710"
    latest_question: str = ""
    case_summary: str = ""
    evidence_subject: str = ""
    evidence_amount: str = ""


class CaseSubmitRequest(BaseModel):
    user_name: str = "王某某"
    phone: str = "133 3107 4710"
    case_summary: str = ""
    evidence_subject: str = ""
    evidence_amount: str = ""


class CaseExtractRequest(BaseModel):
    text: str
    user_name: str = "王某某"
    phone: str = "133 3107 4710"


def safe_filename(title: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|\r\n\t]+", "_", title).strip(" ._")
    return (cleaned or "护薪法律文书")[:80]


def set_document_fonts(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = "SimSun"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def resolve_alignment(tag, inherited_alignment=None):
    classes = set(tag.get("class", []) if hasattr(tag, "get") else [])
    if "text-center" in classes:
        return WD_ALIGN_PARAGRAPH.CENTER
    if "text-right" in classes:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return inherited_alignment


def add_runs_from_node(paragraph, node, inherited_bold=False):
    if isinstance(node, NavigableString):
        text = str(node).replace("\xa0", " ")
        if text:
            run = paragraph.add_run(text)
            run.bold = inherited_bold
        return

    if not getattr(node, "name", None):
        return

    if node.name == "br":
        paragraph.add_run("\n")
        return

    classes = set(node.get("class", []))
    is_bold = inherited_bold or node.name in {"strong", "b"} or "font-bold" in classes
    for child in node.children:
        add_runs_from_node(paragraph, child, is_bold)


def add_html_paragraph(document: Document, tag, alignment=None):
    paragraph = document.add_paragraph()
    classes = set(tag.get("class", []))

    if tag.name == "h1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for child in tag.children:
            add_runs_from_node(paragraph, child, True)
        for run in paragraph.runs:
            run.font.size = Pt(18)
        return

    if tag.name in {"h2", "h3"}:
        paragraph.alignment = alignment
        for child in tag.children:
            add_runs_from_node(paragraph, child, True)
        for run in paragraph.runs:
            run.font.size = Pt(14)
        return

    paragraph.alignment = alignment
    if "indent-8" in classes:
        paragraph.paragraph_format.first_line_indent = Pt(24)

    for child in tag.children:
        add_runs_from_node(paragraph, child)


def append_html_blocks(document: Document, container, inherited_alignment=None):
    block_tags = {"h1", "h2", "h3", "p", "li"}
    for child in container.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                document.add_paragraph(text)
            continue

        if not getattr(child, "name", None) or child.name in {"script", "style"}:
            continue

        alignment = resolve_alignment(child, inherited_alignment)
        if child.name in block_tags:
            add_html_paragraph(document, child, alignment)
        elif child.name == "br":
            document.add_paragraph()
        else:
            append_html_blocks(document, child, alignment)


def build_docx_bytes(title: str, html: str) -> bytes:
    document = Document()
    set_document_fonts(document)

    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    soup = BeautifulSoup(html, "html.parser")
    root = soup.body or soup
    append_html_blocks(document, root)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def read_jsonl(path: Path) -> list[dict]:
    if not path.exists():
        return []

    rows = []
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                rows.append(json.loads(line))
            except json.JSONDecodeError:
                continue
    return rows


def parse_timestamp(value: str) -> datetime | None:
    if not value:
        return None
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S"):
        try:
            return datetime.strptime(value[:19], fmt)
        except ValueError:
            continue
    return None


def make_request_id(prefix: str = "HX") -> str:
    return f"{prefix}{datetime.now().strftime('%Y%m%d%H%M%S%f')[:17]}"


def phase_label(phase: str) -> str:
    labels = {
        "phase2": "第二阶段人工援助",
        "phase4": "第四阶段承办联系",
        "doc": "第三阶段文书导出",
        "submit": "第四阶段提交预审",
    }
    return labels.get(phase, phase or "平台记录")


CN_DIGITS = {"零": 0, "〇": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
CN_UNITS = {"十": 10, "百": 100, "千": 1000}


def parse_chinese_amount_number(text: str) -> int:
    normalized = re.sub(r"[元块人民币整\s]", "", str(text or "").replace("〇", "零"))
    if not normalized:
        return 0

    short_wan = re.fullmatch(r"([一二两三四五六七八九十百千]+)万([一二两三四五六七八九])", normalized)
    if short_wan:
        return parse_chinese_amount_number(short_wan.group(1)) * 10000 + CN_DIGITS[short_wan.group(2)] * 1000

    short_thousand = re.fullmatch(r"([一二两三四五六七八九])千([一二两三四五六七八九])", normalized)
    if short_thousand:
        return CN_DIGITS[short_thousand.group(1)] * 1000 + CN_DIGITS[short_thousand.group(2)] * 100

    short_hundred = re.fullmatch(r"([一二两三四五六七八九])百([一二两三四五六七八九])", normalized)
    if short_hundred:
        return CN_DIGITS[short_hundred.group(1)] * 100 + CN_DIGITS[short_hundred.group(2)] * 10

    total = 0
    section_text = normalized
    if "万" in normalized:
        wan_part, section_text = normalized.split("万", 1)
        total += parse_chinese_amount_number(wan_part) * 10000

    section = 0
    current_number = 0
    for ch in section_text:
        if ch in CN_DIGITS:
            current_number = CN_DIGITS[ch]
        elif ch in CN_UNITS:
            section += (current_number or 1) * CN_UNITS[ch]
            current_number = 0
    return total + section + current_number


def format_amount(amount: int | None) -> str:
    return f"¥ {amount:,}" if amount else "待补充"


def extract_amount_yuan(text: str) -> int | None:
    text = text or ""
    chinese_money = re.search(r"([一二两三四五六七八九十百千万零〇]{2,12})\s*(?:元|块|人民币|整)?", text)
    if chinese_money and re.search(r"[万千百十]", chinese_money.group(1)):
        amount = parse_chinese_amount_number(chinese_money.group(1))
        if amount >= 1000:
            return amount

    digit_wan = re.search(r"(\d+(?:\.\d+)?)\s*万(?:\s*(\d{1,4}))?", text)
    if digit_wan:
        return int(float(digit_wan.group(1)) * 10000 + int(digit_wan.group(2) or 0))

    digit_yuan = re.search(r"(\d{1,3}(?:,\d{3})*|\d+)(?:\.\d{1,2})?\s*[元块]", text)
    if digit_yuan:
        return int(digit_yuan.group(1).replace(",", ""))
    return None


def normalize_subject_candidate(candidate: str) -> str:
    value = str(candidate or "")
    value = re.sub(r"[“”\"『』《》]", "", value)
    value = re.sub(r"^(?:您说的|所谓的|这个|那个|该|欠薪主体|用工主体|用人单位|被申请人|被告|雇主|老板)\s*(?:是|为)?\s*", "", value)
    value = re.sub(r"(?:全称|完整名称|身份信息|统一社会信用代码|联系方式|电话|地址|盖章|签字|签名|是什么|是否|需要|请|吗|呢|？|\?).*$", "", value)
    value = re.split(r"[，。；、,\n\r]", value)[0].strip()

    org_match = re.search(r"[\u4e00-\u9fa5A-Za-z0-9（）()·]{2,40}?(?:公司|项目部|工程部|分包商|劳务队|班组)", value)
    if org_match:
        value = org_match.group(0)

    value = re.sub(r"^(?:在|到|给|跟|为)", "", value)
    value = re.sub(r"(?:那个|这个)?(?:工地|项目|现场)(?:干活|干木工|务工|上班|做工)?.*$", "", value)
    value = re.sub(r"(?:干活|干木工|务工|上班|做工).*$", "", value)
    return value.strip()


def is_useful_subject(candidate: str) -> bool:
    if not candidate or len(candidate) < 2 or len(candidate) > 40:
        return False
    if re.fullmatch(r"(?:公司|单位|雇主|老板|被告|被申请人|项目|工地|包工头|劳动者|申请人)", candidate):
        return False
    if re.search(r"(?:某|XX|xxx|未知|待补充|不清楚|全称|姓名|身份|信息|电话|地址|证据|欠条|工资|金额|劳动合同|工牌)", candidate, re.I):
        return False
    return bool(re.search(r"(?:公司|项目部|工程部|分包商|劳务队|班组)$|^[\u4e00-\u9fa5A-Za-z0-9·]{2,8}(?:老板|包工头)$", candidate))


def extract_debtor_subject(text: str) -> str:
    patterns = [
        r"(?:在|到|给|跟|为|受雇于|入职|就职于)\s*[“\"]?([\u4e00-\u9fa5A-Za-z0-9（）()·]{2,40}?(?:公司|项目部|工程部|分包商|劳务队|班组))[”\"]?(?=那个|这个|的|工地|项目|干|做|上班|务工|施工|$)",
        r"(?:欠薪主体|用工主体|用人单位|被申请人|被告|雇主|劳务公司)\s*[:：]\s*[“\"]?([^，。；\n\r]{2,50})",
        r"[“\"]([\u4e00-\u9fa5A-Za-z0-9（）()·]{2,40}?(?:公司|项目部|工程部|分包商|劳务队|班组))[”\"]",
        r"([\u4e00-\u9fa5A-Za-z0-9（）()·]{2,40}?(?:公司|项目部|工程部|分包商|劳务队|班组))",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text or ""):
            candidate = normalize_subject_candidate(match.group(1))
            if is_useful_subject(candidate):
                return candidate

    person_match = re.search(r"(?:老板|包工头|班组长)\s*([一-龥]{2,4})", text or "")
    if person_match:
        return f"{person_match.group(1)}老板"
    return "待补充"


def extract_project_site(text: str) -> str:
    patterns = [
        r"(?:在|到)\s*([^，。；\n\r]{2,40}?(?:工地|项目|工程|现场))",
        r"(?:项目名称|工程名称|施工地点|工地位置)\s*[:：]\s*([^，。；\n\r]{2,60})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text or "")
        if match:
            value = re.sub(r"(?:干活|干木工|务工|上班|做工).*$", "", match.group(1)).strip()
            if 2 <= len(value) <= 60:
                return value
    return "待补充"


def resolve_relative_year(prefix: str | None) -> int:
    current_year = datetime.now().year
    if prefix == "去年":
        return current_year - 1
    if prefix == "明年":
        return current_year + 1
    return current_year


def extract_timeline(text: str) -> list[dict]:
    timeline = []
    for match in re.finditer(r"(去年|今年|明年)?\s*(\d{1,2})\s*月(?:份)?", text or ""):
        month = int(match.group(2))
        if not 1 <= month <= 12:
            continue
        year = resolve_relative_year(match.group(1))
        nearby = text[max(0, match.start() - 16):match.end() + 18]
        if re.search(r"(?:入场|开始|上班|务工|进场|施工)", nearby):
            title = "开始务工"
        elif re.search(r"(?:完工|离场|结束|干到|停工)", nearby):
            title = "工程完工或停止务工"
        elif re.search(r"(?:干|做)", nearby):
            title = "务工时间节点"
        else:
            title = "案情时间节点"
        timeline.append({"date": f"{year}年{month}月", "title": title, "detail": nearby.strip()})

    if re.search(r"(?:欠条|结算单|工资单)", text or ""):
        timeline.append({"date": "当前", "title": "已持有书面证据", "detail": "案情中提到欠条、结算单或工资单等书面材料。"})
    if re.search(r"(?:跑了|失联|联系不上|拖欠|不给|没给|差我)", text or ""):
        timeline.append({"date": "当前", "title": "发生欠薪争议", "detail": "案情中提到欠薪、失联或拒付工资。"})

    deduped = []
    seen = set()
    for item in timeline:
        key = (item["date"], item["title"])
        if key not in seen:
            seen.add(key)
            deduped.append(item)
    return deduped[:6]


EVIDENCE_KEYWORDS = {
    "欠条": ["欠条", "借条"],
    "工资结算材料": ["结算单", "工资单", "工资表", "工资条", "工程量确认单"],
    "聊天记录": ["微信", "聊天记录", "短信", "施工群", "群聊"],
    "转账记录": ["转账", "银行流水", "收款", "付款记录"],
    "劳动或用工证明": ["劳动合同", "协议", "工牌", "考勤", "打卡", "工地照片", "工作服"],
    "证人证言": ["工友", "证人", "班组", "证明"],
    "录音录像": ["录音", "录像", "视频"],
}


def extract_evidence_items(text: str) -> list[str]:
    items = []
    for label, keywords in EVIDENCE_KEYWORDS.items():
        if any(keyword in (text or "") for keyword in keywords):
            items.append(label)
    return items


def build_missing_items(profile: dict) -> list[str]:
    missing = []
    if profile["debtor_subject"] == "待补充":
        missing.append("欠薪主体身份信息：公司全称、统一社会信用代码，或包工头姓名、电话、身份证线索。")
    if not profile["amount_yuan"]:
        missing.append("欠薪金额依据：欠条、工资结算单、聊天确认、转账记录或手写明细。")
    if profile["work_period"] == "待补充":
        missing.append("务工时间段：入场时间、完工/离场时间，以及期间实际出勤情况。")
    if profile["project_site"] == "待补充":
        missing.append("项目地点信息：工地名称、项目地址、总包或分包单位线索。")
    if not profile["evidence_items"]:
        missing.append("基础证据材料：欠条、聊天记录、工友证明、工地照片、考勤或工牌。")
    elif "欠条" not in profile["evidence_items"] and "工资结算材料" not in profile["evidence_items"]:
        missing.append("书面结算证据：优先补欠条、结算单，或让对方在聊天中确认欠款金额。")
    return missing


def build_next_actions(profile: dict) -> list[str]:
    actions = []
    if profile["debtor_subject"] == "待补充":
        actions.append("先补齐欠薪主体，至少明确公司全称或包工头姓名和联系方式。")
    if not profile["amount_yuan"]:
        actions.append("把欠薪金额写清楚，最好用欠条、结算单或聊天记录固定。")
    if profile["evidence_items"]:
        actions.append("保留证据原件和截图原图，按时间顺序整理成证据目录。")
    actions.append("准备身份证明、联系电话、工地地点和工友联系方式，便于投诉、起诉或申请支持起诉。")
    return actions[:4]


def build_tracking_plan(profile: dict) -> dict:
    if profile.get("evidence_status", "").startswith("证据较充分"):
        current = "材料初审中"
        next_step = "承办人员核对主体、金额和证据原件后，可进入支持起诉审查。"
        expected = "1 个工作日内反馈初审意见"
    elif profile.get("evidence_items"):
        current = "待补充关键材料"
        next_step = "优先补齐缺失材料，补充后再提交承办人员复核。"
        expected = "补充材料后 1 个工作日内复核"
    else:
        current = "线索登记中"
        next_step = "先补充欠薪主体、金额依据和至少一种基础证据。"
        expected = "材料补齐后进入初审"

    return {
        "current_status": current,
        "expected_response": expected,
        "next_step": next_step,
        "handler": "民事检察部门 / 法律援助联络员",
        "anxiety_note": "系统已登记线索；请先保留原件和截图，不需要重复讲述案情，后续按补强清单逐项推进。",
        "steps": [
            {"name": "线索登记", "status": "done", "detail": "已记录咨询内容、联系方式和初步案情。"},
            {"name": "材料初审", "status": "active", "detail": current},
            {"name": "补强证据", "status": "pending", "detail": "根据缺失清单补充主体、金额、时间和证据。"},
            {"name": "支持起诉审查", "status": "pending", "detail": "审查是否符合支持起诉条件并形成意见。"},
        ],
    }


def build_case_profile(text: str, user_name: str = "王某某", phone: str = "133 3107 4710") -> dict:
    text = text or ""
    amount = extract_amount_yuan(text)
    timeline = extract_timeline(text)
    dated_nodes = [item for item in timeline if re.match(r"\d{4}年\d{1,2}月", item["date"])]
    work_period = "待补充"
    if len(dated_nodes) >= 2:
        work_period = f"{dated_nodes[0]['date']} 至 {dated_nodes[1]['date']}"
    elif dated_nodes:
        work_period = dated_nodes[0]["date"]

    evidence_items = extract_evidence_items(text)
    profile = {
        "worker_name": user_name or "待补充",
        "phone": phone or "",
        "debtor_subject": extract_debtor_subject(text),
        "amount_yuan": amount,
        "amount_display": format_amount(amount),
        "work_period": work_period,
        "project_site": extract_project_site(text),
        "evidence_items": evidence_items,
        "timeline": timeline,
    }

    score = 0
    score += 2 if profile["debtor_subject"] != "待补充" else 0
    score += 2 if profile["amount_yuan"] else 0
    score += 1 if profile["work_period"] != "待补充" else 0
    score += 1 if profile["project_site"] != "待补充" else 0
    score += min(2, len(evidence_items))

    if score >= 7:
        evidence_status = "证据较充分，可进入文书生成与预审"
    elif score >= 4:
        evidence_status = "已有初步证据，建议补强关键材料"
    else:
        evidence_status = "证据不足，需补充"

    profile["missing_items"] = build_missing_items(profile)
    profile["next_actions"] = build_next_actions(profile)
    profile["risk_flags"] = [item for item in [
        "欠薪主体不明确" if profile["debtor_subject"] == "待补充" else "",
        "金额缺少稳定依据" if not profile["amount_yuan"] else "",
        "务工时间不完整" if profile["work_period"] == "待补充" else "",
        "书面证据不足" if not evidence_items else "",
    ] if item]
    profile["evidence_status"] = evidence_status
    profile["confidence"] = round(min(score, 8) / 8, 2)
    profile["tracking_plan"] = build_tracking_plan(profile)
    return profile


def record_id(*parts: str) -> str:
    raw = "|".join(str(part or "") for part in parts)
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()[:12]


def engine_error_message(error: Exception) -> str:
    raw = str(error)
    if "401" in raw or "Authentication Fails" in raw or "invalid" in raw.lower() and "api key" in raw.lower():
        return "远程模型认证失败，请检查后端 Code/.env 中的 DEEPSEEK_API_KEY 是否为有效 DeepSeek 密钥。"
    return "远程模型暂时不可用，已切换为本地应急指引。"

def build_local_fallback_answer(query: str, sources: list) -> str:
    source_names = [item.get("filename", "本地案例库") for item in sources[:2]]
    source_text = "、".join(source_names) if source_names else "本地劳动报酬纠纷知识库"
    profile = build_case_profile(query)
    missing_text = "\n".join(f"- {item}" for item in profile["missing_items"][:4]) or "- 目前关键要素较完整，建议继续保留原始证据。"
    evidence_text = "、".join(profile["evidence_items"]) or "暂未识别到明确证据材料"
    next_action_text = "\n".join(f"{idx}. {item}" for idx, item in enumerate(profile["next_actions"], start=1))
    return f"""我先给您一个可立即执行的维权方案。

**初步识别**

您描述的是一起追索劳动报酬纠纷。我先按现有信息抽取如下：欠薪主体为“{profile['debtor_subject']}”，欠薪金额为“{profile['amount_display']}”，务工时间为“{profile['work_period']}”，项目地点为“{profile['project_site']}”，已识别证据为“{evidence_text}”。

**证据链状态**

{profile['evidence_status']}。当前最需要补强的是：

{missing_text}

**下一步建议**

{next_action_text}

如果您是农民工、取证困难、自己起诉能力弱，可以向检察机关申请支持起诉，请求帮助固定证据、梳理被告主体和诉讼请求。

我已参考 {source_text} 做初步研判。远程智能模型当前不可用时，上面是本地应急指引；等密钥恢复后，系统会继续生成更完整的支持起诉分析和文书要点。"""

# ================= 3. Core API Endpoints =================
@app.get("/")
async def serve_frontend():
    return FileResponse(CODE_DIR / "Web" / "index.html")


@app.post("/api/case-extract")
async def extract_case(payload: CaseExtractRequest):
    profile = build_case_profile(payload.text, payload.user_name, payload.phone)
    backend_log(
        "CASE_EXTRACT",
        f"{payload.user_name}({payload.phone}) | subject={profile['debtor_subject']} | amount={profile['amount_display']} | status={profile['evidence_status']}"
    )
    return {"case": profile}


@app.post("/api/chat")
async def chat_endpoint(request: ChatRequest):
    legal_context = "No relevant legal documents found in the local database."
    source_items = []

    selected_model = "deepseek-reasoner" if request.mode == "prism" else "deepseek-chat"
    
    backend_log(
        "PHASE2 AI_CHAT_REQUEST",
        f"{request.user_name}({request.phone}) | {request.query[:300]} | Engine: {request.mode.upper()} ({selected_model})"
    )
    
    # Keywords specific to labor disputes and wage protection
    legal_keywords = [
        '法律', '法条', '条款', '民法', '刑法', '公司法', '劳动法', '合同法', '合规', '法',
        '章程', '准则', '条例', '规定', '司法解释', '原告', '被告', '第三人', 
        '连带责任', '债权', '债务', '担保', '诉讼', '仲裁', '起诉', '申诉', '判决', '调解', 
        '证据', '举证', '违约', '赔偿', '效力', '判例', '案例', '协议', '合同', '起诉状',
        '欠薪', '工资', '薪水', '薪资', '农民工', '包工头', '劳动关系', '支持起诉', '援助', '维权'
    ]
    
    is_legal_query = any(k in request.query for k in legal_keywords) or len(request.query) > 10

    if not is_legal_query:
        print("[Intent Filter] Daily chat detected, skipping RAG retrieval...")
    else:
        print("[Intent Filter] Legal query detected, initiating ChromaDB retrieval...")
        raw_results = vectordb.similarity_search_with_score(request.query, k=request.top_k)
        legal_context = ""
        source_items = []
        
        for i, (doc, score) in enumerate(raw_results):
            if score < request.score_threshold:
                filename = doc.metadata.get('source', 'Unknown File')
                legal_context += f"\n--- Document {i+1} (Source: {filename}) ---\n{doc.page_content}\n"
                source_items.append({
                    "filename": filename, 
                    "score": round(score, 4),
                    "content_preview": doc.page_content[:30] + "..."
                })

    case_profile = build_case_profile(request.query, request.user_name, request.phone)
    case_profile_text = json.dumps(case_profile, ensure_ascii=False, indent=2)

# === 增强版 System Prompt：深度结合 RAG 与 2026 法律背景 ===
    final_system_prompt = f"""您是“护薪”检察支持起诉智能平台的智能助理 Justitia（也可以自称“小朱”），由 Huang Zitong 开发，专门服务于北京市西城区人民检察院。您的核心使命是协助农民工追索劳动报酬，并辅助检察官进行“支持起诉”的案件预审。注意，你的服务对象是维权的农民工群体，因此请保持语言精密而不失通俗，专业而不失关怀。

    [系统指令深度对齐]：
    1. 时间锚点：当前为 2026 年春季，请确保所有建议符合最新的法律时效。
    2. 双重身份：您既是农民工的贴心法律向导，又是检察官的专业审查助理。对劳动者请使用通俗、温暖、感性的语言；对案件分析请保持严密的法理逻辑。
    3. RAG 核心驱动：
       - 【法条库】：基于检索到的 [Local Legal Context]，优先引用内置的最新司法解释与《保障农民工工资支付条例》。
       - 【文书模板】：当用户信息基本完整时，必须引导并参考知识库中的“起诉状”、“支持起诉申请书”等标准模板格式生成预览。

    [执行指南 - 核心五要素提取]：
    您必须从 OCR 提取的文本或用户对话中精准锁定：
    1) 欠薪主体（用人单位全称/包工头姓名/项目部名称）。
    2) 劳动者身份。
    3) 确切的欠薪数额（需与证据中的数字对齐）。
    4) 务工时间段及项目名称。
    5) 现有证据清单（欠条、结算单、微信记录、工牌等）。

    [法律研判逻辑]：
    - 证据校验：如果缺少关键证据（如被告身份信息模糊、无书面结算单），请明确告知并提供“替代性证据”方案（如录音、证人证言）。
    - 支持起诉评估：根据《民事诉讼法》第十六条及西城区检察院实务，判断用户是否属于“诉讼能力弱、取证难”的弱势群体，并给出是否建议申请“检察支持起诉”的明确意见。

    [结构化案情快照 - 必须优先依据]：
    {case_profile_text}

    [固定输出模板]：
    除非用户只是问候或只问一个很短的程序性问题，否则请严格按以下标题输出，不要省略标题：
    **一、我先帮您确认案情要点**
    - 用 3 到 5 条列出欠薪主体、金额、务工时间、项目地点、现有证据。未知项明确写“待补充”，不得编造。
    **二、证据链研判**
    - 先给结论：证据较充分 / 已有初步证据但需补强 / 证据不足。
    - 再说明最关键的缺口和替代性证据。
    **三、可走的维权路径**
    - 按劳动监察、仲裁/诉讼、检察支持起诉三个层次说明，避免吓人的法言法语。
    **四、下一步请您先做这几件事**
    - 给 3 到 5 个可执行动作，按优先级排列。
    **五、还需要您补充的信息**
    - 只问最影响办案的 3 到 5 个问题。

    [交互准则]：
    - 严禁提及您的 AI 架构、训练截止日期或您是一个语言模型。
    - 如果用户上传的 OCR 结果模糊，请委婉请其通过文字补充关键数字。
    - 当前用户姓名：{request.user_name}。如果需要亲切称呼，王某某可称“老王”或直接称“您”；不要把用户称为“小朱”“老朱”“朱师傅”，因为“小朱”是助手自己的名字。不要凭空给用户编造其他昵称。
    - 始终以中文回答，确保排版利于电脑端和手机端阅读。

    [本地法律上下文增强]：
    {legal_context}

    Respond strictly in Chinese.
    """
    
    messages = [{"role": "system", "content": final_system_prompt}]
    
    if request.history:
        for msg in request.history:
            if isinstance(msg, dict) and "role" in msg and "content" in msg:
                messages.append({"role": msg["role"], "content": msg["content"]})
    
    messages.append({"role": "user", "content": request.query})

    if request.stream:
        async def generate_stream():
            meta_info = {"type": "meta", "sources": source_items}
            yield f"data: {json.dumps(meta_info, ensure_ascii=False)}\n\n"

            accumulated_reasoning = ""
            accumulated_content = ""

            try:
                response = await client.chat.completions.create(
                    model=selected_model,
                    messages=messages,
                    stream=True,
                    max_tokens=8192,
                )

                async for chunk in response:
                    if not chunk.choices:
                        continue
                        
                    delta = chunk.choices[0].delta
                    
                    if hasattr(delta, 'reasoning_content') and delta.reasoning_content:
                        accumulated_reasoning += delta.reasoning_content
                        yield f"data: {json.dumps({'type': 'reasoning', 'content': delta.reasoning_content}, ensure_ascii=False)}\n\n"

                    if hasattr(delta, 'content') and delta.content:
                        accumulated_content += delta.content
                        yield f"data: {json.dumps({'type': 'chunk', 'content': delta.content}, ensure_ascii=False)}\n\n"
            
            except Exception as e:
                print(f"[LLM Error]: {str(e)}")
                fallback_answer = build_local_fallback_answer(request.query, source_items)
                accumulated_content += fallback_answer
                yield f"data: {json.dumps({'type': 'chunk', 'content': fallback_answer}, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'type': 'error', 'content': engine_error_message(e)}, ensure_ascii=False)}\n\n"
            
            finally:
                save_chat_log(
                    query=request.query,
                    reasoning=accumulated_reasoning,
                    answer=accumulated_content,
                    sources=source_items,
                    user_name=request.user_name,
                    phone=request.phone,
                    case_profile=build_case_profile(f"{request.query}\n{accumulated_content}", request.user_name, request.phone),
                )
                
            yield "data: [DONE]\n\n"

        return StreamingResponse(generate_stream(), media_type="text/event-stream")
    
    else:
        try:
            response = await client.chat.completions.create(
                model=selected_model,
                messages=messages,
                max_tokens=8192,
            )

            full_answer = response.choices[0].message.content
            full_reasoning = getattr(response.choices[0].message, 'reasoning_content', "")
            engine_error = None
        except Exception as e:
            print(f"[LLM Error]: {str(e)}")
            full_answer = build_local_fallback_answer(request.query, source_items)
            full_reasoning = ""
            engine_error = engine_error_message(e)

        save_chat_log(
            request.query,
            full_reasoning,
            full_answer,
            source_items,
            request.user_name,
            request.phone,
            case_profile=build_case_profile(f"{request.query}\n{full_answer}", request.user_name, request.phone),
        )

        payload = {"answer": full_answer, "sources": source_items, "reasoning": full_reasoning}
        if engine_error:
            payload["error"] = engine_error
        return payload

@app.post("/api/export-docx")
async def export_docx(payload: DocExportRequest):
    title = safe_filename(payload.title)
    if not payload.html.strip():
        raise HTTPException(status_code=400, detail="文书内容为空，无法导出")

    record_platform_event("document_exported", {
        "stage": phase_label("doc"),
        "user_name": payload.user_name,
        "phone": payload.phone,
        "question": title,
        "status": "已生成文书",
        "summary": f"导出文书：{title}",
    })

    docx_bytes = build_docx_bytes(title, payload.html)
    filename = f"{title}.docx"
    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"
    }
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.post("/api/human-support")
async def request_human_support(payload: HumanSupportRequest):
    request_id = make_request_id("HS")
    assignee = "法律援助志愿者" if payload.phase == "phase2" else "民事检察承办助理"
    status = "待人工跟进"
    summary_parts = [
        payload.case_summary.strip(),
        f"主体：{payload.evidence_subject}" if payload.evidence_subject else "",
        f"金额：{payload.evidence_amount}" if payload.evidence_amount else "",
    ]
    summary = "；".join(part for part in summary_parts if part) or payload.latest_question[:160]

    record_platform_event("human_support_requested", {
        "request_id": request_id,
        "stage": phase_label(payload.phase),
        "user_name": payload.user_name,
        "phone": payload.phone,
        "question": payload.latest_question[:500],
        "summary": summary,
        "status": status,
        "assignee": assignee,
    })

    return {
        "request_id": request_id,
        "status": status,
        "assignee": assignee,
        "message": f"已登记人工协助请求，编号 {request_id}。{assignee}会根据后台记录继续处理。",
    }


@app.post("/api/case-submit")
async def submit_case(payload: CaseSubmitRequest):
    request_id = make_request_id("CS")
    profile = build_case_profile(
        f"{payload.case_summary}\n主体：{payload.evidence_subject}\n金额：{payload.evidence_amount}",
        payload.user_name,
        payload.phone,
    )
    tracking = profile["tracking_plan"]
    record_platform_event("case_submitted", {
        "request_id": request_id,
        "stage": phase_label("submit"),
        "user_name": payload.user_name,
        "phone": payload.phone,
        "question": payload.case_summary[:500],
        "summary": f"主体：{payload.evidence_subject or '待补充'}；金额：{payload.evidence_amount or '待补充'}",
        "status": tracking["current_status"],
        "assignee": tracking["handler"],
        "case_profile": profile,
    })
    return {
        "request_id": request_id,
        "status": tracking["current_status"],
        "tracking_plan": tracking,
        "message": f"预审卷宗已登记，编号 {request_id}。",
    }


@app.get("/api/admin/records")
async def admin_records(days: int = 7):
    cutoff = datetime.now() - timedelta(days=max(1, min(days, 30)))
    records = []

    for item in read_jsonl(LOG_FILE_PATH):
        ts = parse_timestamp(item.get("timestamp", ""))
        if ts is None or ts < cutoff:
            continue
        answer = item.get("justitia_answer", "")
        query = item.get("user_query", "")
        profile = item.get("case_profile") or build_case_profile(f"{query}\n{answer}", item.get("user_name", "王某某"), item.get("phone", "133 3107 4710"))
        records.append({
            "id": record_id(item.get("timestamp"), item.get("phone"), query, "chat"),
            "timestamp": item.get("timestamp"),
            "stage": "第二阶段 AI研判",
            "farmer_name": item.get("user_name", "王某某"),
            "phone": item.get("phone", "133 3107 4710"),
            "question": query,
            "summary": answer[:180],
            "status": item.get("status", "AI 已答复"),
            "assignee": "Justitia 护薪助手",
            "answer": answer,
            "reasoning": item.get("justitia_thought", ""),
            "sources": item.get("reference_sources", []),
            "case_profile": profile,
        })

    for item in read_jsonl(EVENT_LOG_PATH):
        ts = parse_timestamp(item.get("timestamp", ""))
        if ts is None or ts < cutoff:
            continue
        question = item.get("question", "")
        summary = item.get("summary", "")
        profile = item.get("case_profile") or build_case_profile(f"{question}\n{summary}", item.get("user_name", "王某某"), item.get("phone", "133 3107 4710"))
        records.append({
            "id": record_id(item.get("timestamp"), item.get("phone"), item.get("request_id"), item.get("event_type")),
            "timestamp": item.get("timestamp"),
            "stage": item.get("stage", item.get("event_type", "平台记录")),
            "farmer_name": item.get("user_name", "王某某"),
            "phone": item.get("phone", "133 3107 4710"),
            "question": question,
            "summary": summary,
            "status": item.get("status", "已记录"),
            "assignee": item.get("assignee", ""),
            "request_id": item.get("request_id", ""),
            "event_type": item.get("event_type", ""),
            "case_profile": profile,
        })

    records.sort(key=lambda row: row.get("timestamp", ""), reverse=True)
    backend_log("ADMIN_DASHBOARD_QUERY", f"days={days} | records={len(records)}")
    return {"days": days, "records": records[:200]}


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        text = ""
        original_filename = file.filename or "uploaded_file"
        filename = original_filename.lower()
        ocr_confidence = None
        backend_log("PHASE2 FILE_UPLOAD", f"{original_filename} | {len(contents)} bytes")
        
        # 1. 处理 PDF/DOCX/TXT (保持原样)
        if filename.endswith('.pdf'):
            pdf_reader = PdfReader(io.BytesIO(contents))
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
        elif filename.endswith('.docx'):
            text = docx2txt.process(io.BytesIO(contents))
        elif filename.endswith('.txt'):
            text = contents.decode('utf-8')
            
        # 2. 处理图片证据 (OCR)
        elif filename.endswith(('.png', '.jpg', '.jpeg', '.bmp')):
            if ocr_engine is None:
                return {"error": f"OCR 引擎未就绪：{ocr_init_error or '请检查 rapidocr/easyocr 与模型文件是否安装完整'}"}

            backend_log("PHASE2 OCR_START", f"{ocr_engine_name} | {original_filename}")
            image = Image.open(io.BytesIO(contents)).convert('RGB')
            image_np = np.array(image)
            text, ocr_confidence = run_ocr(image_np)

            if not text.strip():
                return {"error": "OCR 未识别到有效文字，请换一张更清晰的图片，或直接输入欠条上的金额、签名和日期。"}

            print("\n" + "="*30 + " 扫描结果可视化 " + "="*30, flush=True)
            print(text, flush=True) # 这里会在后端控制台完整输出图片文字
            print("="*76 + "\n", flush=True)
            
            confidence_log = f"，平均置信度: {ocr_confidence}" if ocr_confidence is not None else ""
            backend_log("PHASE2 OCR_DONE", f"{original_filename} | 提取字数: {len(text)}{confidence_log}")
            
        else:
            return {"error": "暂不支持该文件格式"}

        return {
            "filename": original_filename,
            "content_preview": text[:500],
            "full_content": text,
            "ocr_engine": ocr_engine_name if filename.endswith(('.png', '.jpg', '.jpeg', '.bmp')) else None,
            "ocr_confidence": ocr_confidence,
        }
    except Exception as e:
        backend_log("Parse Error", str(e))
        return {"error": f"文件解析失败: {str(e)}"}

if __name__ == "__main__":
    backend_log("SERVER_START", f"Justitia Shield API Server starting at http://{SERVER_HOST}:{SERVER_PORT}")
    uvicorn.run(app, host=SERVER_HOST, port=SERVER_PORT)
